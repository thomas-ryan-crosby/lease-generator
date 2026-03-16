"""
Lease Generator Tool for Crosby Development Company, LLC
Generates commercial lease agreements and amendments as .docx files.
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta
import os
import re
import json

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from num2words import num2words

BLUE = RGBColor(0, 0, 180)
DEFAULT_SAVE_DIR = os.path.join(os.path.expanduser("~"), "OneDrive",
                                "Desktop", "Auto Generated Leases")
CONFIG_FILE = os.path.join(os.path.expanduser("~"), ".lease_generator_config.json")


def load_config():
    """Load saved settings from config file."""
    try:
        with open(CONFIG_FILE, "r") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return {}


def save_config(cfg):
    """Persist settings to config file."""
    with open(CONFIG_FILE, "w") as f:
        json.dump(cfg, f, indent=2)


def get_save_dir():
    """Return the configured save directory, falling back to the default."""
    cfg = load_config()
    return cfg.get("save_dir", DEFAULT_SAVE_DIR)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def get_next_lease_id(doc_type="L"):
    """Get the next unique lease ID in format CDC-X-YYYY-NNNN."""
    save_dir = get_save_dir()
    os.makedirs(save_dir, exist_ok=True)
    counter_file = os.path.join(save_dir, "lease_counter.json")
    year = datetime.today().year
    try:
        with open(counter_file, "r") as f:
            data = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        data = {}
    year_key = str(year)
    current = data.get(year_key, 0)
    current += 1
    data[year_key] = current
    with open(counter_file, "w") as f:
        json.dump(data, f)
    return f"CDC-{doc_type}-{year}-{current:04d}"


def dollars_to_words(amount_str):
    """Convert a dollar amount string like '862.00' to written form."""
    amount_str = amount_str.replace(",", "").replace("$", "").strip()
    try:
        amount = float(amount_str)
    except ValueError:
        return amount_str
    dollars = int(amount)
    cents = round((amount - dollars) * 100)
    dollar_words = num2words(dollars, lang="en").replace(",", "")
    # Capitalize first letter, rest lowercase
    dollar_words = dollar_words[0].upper() + dollar_words[1:]
    if cents == 0:
        return f"{dollar_words} dollars ({format_currency(amount_str)})"
    else:
        cents_words = num2words(cents, lang="en")
        return f"{dollar_words} and {cents_words} cents ({format_currency(amount_str)})"


def number_words(n):
    """Convert a number to lowercase words, e.g. 90 -> 'ninety', 3.5 -> 'three and one half'."""
    n = float(n)
    if n == int(n):
        return num2words(int(n), lang="en")
    whole = int(n)
    frac = round(n - whole, 10)
    whole_words = num2words(whole, lang="en")
    # Express fractional part as a simple fraction (e.g. 0.5 -> "one half", 0.25 -> "one quarter")
    frac_map = {0.25: "one quarter", 0.5: "one half", 0.75: "three quarters"}
    frac_words = frac_map.get(frac, num2words(frac, lang="en"))
    return f"{whole_words} and {frac_words}"


def format_currency(amount_str):
    """Format a number string as $X,XXX.XX"""
    amount_str = amount_str.replace(",", "").replace("$", "").strip()
    try:
        amount = float(amount_str)
    except ValueError:
        return amount_str
    return f"${amount:,.2f}"


def format_date_long(date_obj):
    """Format date as 'Month Day, Year' e.g. 'January 1, 2026'"""
    return date_obj.strftime("%B ") + str(date_obj.day) + ", " + str(date_obj.year)


def parse_date(date_str):
    """Parse MM/DD/YYYY to datetime."""
    for fmt in ("%m/%d/%Y", "%m-%d-%Y", "%Y-%m-%d"):
        try:
            return datetime.strptime(date_str.strip(), fmt)
        except ValueError:
            continue
    raise ValueError(f"Cannot parse date: {date_str}")


# ---------------------------------------------------------------------------
# Document builders with blue-text support
#
# Use <<blue>> and <<end>> markers in text strings to indicate user input
# that should be rendered in blue.  Everything else renders in black.
# ---------------------------------------------------------------------------

def _apply_run_fmt(run, font_size=12, font_name="Times New Roman",
                   bold=False, italic=False, color=None):
    """Apply standard formatting to a run."""
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_mixed_runs(paragraph, text, font_size=12, font_name="Times New Roman",
                   bold=False, italic=False):
    """Add runs to a paragraph, coloring <<blue>>...<<end>> segments in blue."""
    parts = re.split(r'(<<blue>>|<<end>>)', text)
    in_blue = False
    for part in parts:
        if part == '<<blue>>':
            in_blue = True
            continue
        if part == '<<end>>':
            in_blue = False
            continue
        if not part:
            continue
        run = paragraph.add_run(part)
        color = BLUE if in_blue else None
        _apply_run_fmt(run, font_size, font_name, bold, italic, color)


def add_paragraph(doc, text, bold=False, alignment=None, font_size=12,
                  space_before=0, space_after=0, font_name="Times New Roman",
                  first_line_indent=None, italic=False, color=None):
    """Add a paragraph with formatting.  Supports <<blue>>...<<end>> markers."""
    p = doc.add_paragraph()
    if '<<blue>>' in text:
        add_mixed_runs(p, text, font_size, font_name, bold, italic)
    else:
        run = p.add_run(text)
        _apply_run_fmt(run, font_size, font_name, bold, italic, color)
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent is not None:
        pf.first_line_indent = Inches(first_line_indent)
    return p


def add_clause(doc, title, body, font_size=12):
    """Add a clause: bold title followed by body text, matching original format."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(2)

    # Title in bold
    run_title = p.add_run(f"{title}  ")
    _apply_run_fmt(run_title, font_size, bold=True)

    # Body with blue markers
    add_mixed_runs(p, body, font_size)
    return p


def add_sub_items(doc, items, font_size=12):
    """Add indented sub-items."""
    for item in items:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Inches(0.75)
        pf.space_before = Pt(1)
        pf.space_after = Pt(1)
        run = p.add_run(item)
        _apply_run_fmt(run, font_size)


def B(val):
    """Wrap a value in blue markers."""
    return f"<<blue>>{val}<<end>>"


def _remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'none')
        elem.set(qn('w:sz'), '0')
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), 'auto')
        borders.append(elem)
    tblPr.append(borders)


def _sig_cell(cell, text, bold=False, color=None, font_size=12):
    """Format a signature block table cell."""
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    _apply_run_fmt(run, font_size=font_size, bold=bold, color=color)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


# ---------------------------------------------------------------------------
# Lease generation
# ---------------------------------------------------------------------------

def generate_lease(fields):
    """Generate a full commercial lease .docx from the field dictionary."""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    lease_id = get_next_lease_id("L")

    # --- Parse all fields ---
    building_num = fields["building_num"]
    suite = fields["suite"]
    floor = fields["floor"]
    sqft = fields["sqft"]
    start_date = parse_date(fields["start_date"])
    end_date = parse_date(fields["end_date"])
    term_months = fields["term_months"]
    renewal_months = fields["renewal_months"]
    notice_days = fields["notice_days"]
    rent = fields["rent"]
    rent_words = dollars_to_words(rent)
    escalation = fields["escalation"]
    first_payment = parse_date(fields["first_payment_date"])
    utilities = fields["utilities"]
    common_area = fields["common_area"]
    use_of_premises = fields["use_of_premises"]
    buildout = fields["buildout"]
    lobby_name = fields["lobby_name"]
    security_deposit = fields["security_deposit"]
    insurance_amount = fields["insurance_amount"]
    insurance_formatted = format_currency(insurance_amount)
    notice_words = number_words(notice_days)
    escalation_words = number_words(escalation)

    # --- SUMMARY PAGE (Non-Binding) ---
    add_paragraph(doc, "LEASE SUMMARY", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=16,
                  space_before=0, space_after=4)

    add_paragraph(doc,
        "This Lease Summary is provided for informational purposes only and does not "
        "constitute a legally binding agreement.  In the event of any discrepancy between "
        "this summary and the attached Lease Agreement, the terms of the Lease Agreement "
        "shall prevail.",
        font_size=10, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0, space_after=16)

    def add_summary_item(label, value):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.25)
        run = p.add_run("\u2022  ")
        _apply_run_fmt(run, font_size=11)
        run2 = p.add_run(f"{label}:  ")
        _apply_run_fmt(run2, font_size=11, bold=True)
        run3 = p.add_run(value)
        _apply_run_fmt(run3, font_size=11, color=BLUE)

    add_summary_item("Lease ID", lease_id)
    add_summary_item("Lessor",
        "Crosby Development Company, LLC, "
        "#1 Sanctuary Boulevard, Mandeville, LA 70471")
    add_summary_item("Lessee",
        f"{fields['lessee_name']}, {fields['lessee_addr1']}, "
        f"{fields['lessee_addr2']}")
    add_summary_item("Premises",
        f"#{building_num} Sanctuary Blvd, Floor {floor}, Suite {suite} "
        f"({sqft} sqft)")
    add_summary_item("Lease Term",
        f"{term_months} months  ({format_date_long(start_date)} to "
        f"{format_date_long(end_date)})")
    add_summary_item("Monthly Rent", format_currency(rent))
    add_summary_item("Annual Escalation", f"{escalation}%")
    add_summary_item("First Payment Due", format_date_long(first_payment))
    add_summary_item("Renewal",
        f"Auto-renews for {renewal_months}-month terms; "
        f"{notice_days} days written notice to terminate")
    add_summary_item("Utilities", f"Paid by {utilities}")
    if common_area == "None":
        add_summary_item("Common Area Charges", "None")
    else:
        add_summary_item("Common Area Charges", "Paid by Lessee")
    add_summary_item("Use of Premises", use_of_premises)
    sec_dep_summary = security_deposit.strip()
    if sec_dep_summary.lower() == "waived":
        add_summary_item("Security Deposit", "Waived")
    else:
        add_summary_item("Security Deposit", format_currency(sec_dep_summary))
    add_summary_item("Insurance", f"{insurance_formatted} liability per occurrence")
    add_summary_item("Lobby Directory", lobby_name)

    # Page break before metrics
    doc.add_page_break()

    # --- METRICS PAGE (For Lessor Internal Use) ---
    add_paragraph(doc, "LEASE FINANCIAL METRICS", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=16,
                  space_before=0, space_after=2)
    add_paragraph(doc, "For Lessor Internal Use Only",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=11,
                  italic=True, space_before=0, space_after=2)
    add_paragraph(doc,
        "These metrics are for internal evaluation purposes only and do not form "
        "part of the Lease Agreement.  They are not to be shared with the Lessee.",
        font_size=9, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0, space_after=12)

    # --- Calculate metrics ---
    monthly_rent_num = float(rent.replace(",", "").replace("$", ""))
    sqft_num = float(sqft.replace(",", ""))
    term_months_num = int(term_months)
    escalation_num = float(escalation)
    term_years = term_months_num / 12

    # Annual rent per SF
    annual_rent_sf = (monthly_rent_num * 12) / sqft_num if sqft_num > 0 else 0

    # Total Lease Value with escalation
    tlv = 0
    months_left = term_months_num
    yr_idx = 0
    while months_left > 0:
        mo_this_yr = min(12, months_left)
        yr_rent = monthly_rent_num * ((1 + escalation_num / 100) ** yr_idx)
        tlv += yr_rent * mo_this_yr
        months_left -= mo_this_yr
        yr_idx += 1

    avg_monthly = tlv / term_months_num if term_months_num > 0 else 0

    # Final year monthly rent
    last_yr_idx = max(0, (term_months_num - 1) // 12)
    final_yr_monthly = monthly_rent_num * ((1 + escalation_num / 100) ** last_yr_idx)

    # Cumulative escalation
    cumulative_esc = ((1 + escalation_num / 100) ** term_years - 1) * 100

    # Security deposit
    sec_dep_raw = security_deposit.strip()
    if sec_dep_raw.lower() == "waived":
        dep_amount = 0
    else:
        dep_amount = float(sec_dep_raw.replace(",", "").replace("$", ""))
    dep_months = dep_amount / monthly_rent_num if monthly_rent_num > 0 else 0
    dep_pct = (dep_amount / tlv * 100) if tlv > 0 else 0
    unsecured = tlv - dep_amount

    # --- Build metrics rows: (metric, value, assessment) ---
    metrics_rows = []

    # Section: Revenue
    metrics_rows.append(("REVENUE METRICS", "", "", True))

    if annual_rent_sf > 19:
        rent_note = "Above LA suburban office avg ($18\u2013$19/SF)"
    elif annual_rent_sf >= 18:
        rent_note = "At LA suburban office avg ($18\u2013$19/SF)"
    else:
        rent_note = "Below LA suburban office avg ($18\u2013$19/SF)"
    metrics_rows.append(("Annual Rent / SF",
        f"${annual_rent_sf:,.2f}/SF/yr", rent_note, False))

    metrics_rows.append(("Total Lease Value",
        f"${tlv:,.2f}",
        f"Total committed revenue over {term_months} months", False))

    metrics_rows.append(("Average Monthly Rent",
        f"${avg_monthly:,.2f}",
        "Blended monthly average with escalation", False))

    metrics_rows.append(("Final Year Monthly Rent",
        f"${final_yr_monthly:,.2f}",
        f"${final_yr_monthly - monthly_rent_num:+,.2f}/mo vs. starting rent", False))

    metrics_rows.append(("Cumulative Escalation",
        f"{cumulative_esc:.1f}%",
        f"Total rent increase over {term_years:.1f}-year term", False))

    # Section: Security & Risk
    metrics_rows.append(("SECURITY & RISK", "", "", True))

    if sec_dep_raw.lower() == "waived":
        metrics_rows.append(("Security Deposit", "Waived",
            "No default protection", False))
    else:
        if dep_months >= 2:
            dep_note = "Strong coverage"
        elif dep_months >= 1:
            dep_note = "Standard coverage"
        else:
            dep_note = "Below standard \u2014 consider increasing"
        metrics_rows.append(("Security Deposit Coverage",
            f"{dep_months:.1f} months of rent", dep_note, False))

    metrics_rows.append(("Unsecured Lease Value",
        f"${unsecured:,.2f}",
        "Total Lease Value minus deposit", False))

    metrics_rows.append(("Deposit as % of TLV",
        f"{dep_pct:.2f}%",
        "Proportion of total deal that is secured", False))

    # Section: Deal Quality
    metrics_rows.append(("DEAL QUALITY", "", "", True))

    if utilities == "Lessee" and common_area == "Lessee Pays":
        exp_type, exp_note = "NNN-like", "Favorable \u2014 Lessee absorbs operating costs"
    elif utilities == "Lessee" or common_area == "Lessee Pays":
        exp_type, exp_note = "Modified Gross", "Partially favorable \u2014 some cost recovery"
    else:
        exp_type, exp_note = "Full Service Gross", "Lessor absorbs utility and common area costs"
    metrics_rows.append(("Expense Structure", exp_type, exp_note, False))

    if term_months_num < 24:
        term_lbl, term_note = "Short-term", "Higher turnover risk"
    elif term_months_num <= 60:
        term_lbl, term_note = "Medium-term", "Standard for office"
    else:
        term_lbl, term_note = "Long-term", "Strong income stability"
    metrics_rows.append(("Lease Term Strength",
        f"{term_lbl} ({term_months} mo)", term_note, False))

    if escalation_num > 3:
        esc_note = "Above inflation \u2014 purchasing power grows"
    elif escalation_num >= 3:
        esc_note = "At inflation \u2014 purchasing power maintained"
    else:
        esc_note = "Below inflation \u2014 purchasing power may erode"
    metrics_rows.append(("Escalation vs. Inflation",
        f"{escalation}% vs. ~3% CPI", esc_note, False))

    notice_num = int(notice_days)
    renewal_num = int(renewal_months)
    if notice_num >= 90:
        renew_note = f"Strong \u2014 auto-renews {renewal_num} mo, {notice_num}-day notice"
    elif notice_num >= 60:
        renew_note = f"Good \u2014 auto-renews {renewal_num} mo, {notice_num}-day notice"
    else:
        renew_note = f"Moderate \u2014 {notice_num}-day notice may limit lead time"
    metrics_rows.append(("Renewal Strength", "Auto-renew", renew_note, False))

    # --- Render metrics table ---
    mtable = doc.add_table(rows=len(metrics_rows) + 1, cols=3)
    mtable.style = 'Table Grid'
    mtable.allow_autofit = False

    # Column widths
    for r in mtable.rows:
        r.cells[0].width = Inches(1.8)
        r.cells[1].width = Inches(1.7)
        r.cells[2].width = Inches(3.0)

    # Header row
    for ci, header in enumerate(["Metric", "Value", "Assessment"]):
        cell = mtable.cell(0, ci)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(header)
        _apply_run_fmt(run, font_size=10, bold=True)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for ri, (metric, value, assessment, is_section) in enumerate(metrics_rows, start=1):
        if is_section:
            # Section header: bold text in first cell, all cells shaded
            for ci in range(3):
                cell = mtable.cell(ri, ci)
                p = cell.paragraphs[0]
                p.clear()
                if ci == 0:
                    run = p.add_run(metric)
                    _apply_run_fmt(run, font_size=10, bold=True)
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'E8E8E8')
                shading.set(qn('w:val'), 'clear')
                cell._tc.get_or_add_tcPr().append(shading)
        else:
            # Metric name
            p = mtable.cell(ri, 0).paragraphs[0]
            p.clear()
            run = p.add_run(metric)
            _apply_run_fmt(run, font_size=10)

            # Value (bold)
            p = mtable.cell(ri, 1).paragraphs[0]
            p.clear()
            run = p.add_run(value)
            _apply_run_fmt(run, font_size=10, bold=True)

            # Assessment (italic)
            p = mtable.cell(ri, 2).paragraphs[0]
            p.clear()
            run = p.add_run(assessment)
            _apply_run_fmt(run, font_size=9, italic=True)

    # Page break before the lease
    doc.add_page_break()

    # --- LEASE ID ---
    add_paragraph(doc, f"Lease ID: {lease_id}",
                  alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_size=10,
                  space_before=0, space_after=4, italic=True)

    # --- TITLE ---
    add_paragraph(doc, "LEASE OF COMMERCIAL PROPERTY", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14,
                  space_before=0, space_after=12)

    # --- PARTIES - LESSOR ---
    p = doc.add_paragraph()
    run = p.add_run("Parties")
    _apply_run_fmt(run, bold=True)
    run2 = p.add_run(" Crosby Development Company, LLC")
    _apply_run_fmt(run2)

    add_paragraph(doc, "\t\t\t#1 Sanctuary Boulevard", space_after=0)
    add_paragraph(doc, "\t\t\tMandeville, LA  70471", space_after=0)
    add_paragraph(doc, '\t\t(Hereinafter referred to as "Lessor")',
                  space_after=6)
    add_paragraph(doc, "hereby leases to:", space_after=6)

    # --- PARTIES - LESSEE (blue) ---
    add_paragraph(doc, fields["lessee_name"],
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, color=BLUE)
    add_paragraph(doc, fields["lessee_addr1"],
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, color=BLUE)
    add_paragraph(doc, fields["lessee_addr2"],
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, color=BLUE)
    add_paragraph(doc, '(Hereinafter referred to as "Lessee")',
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    # --- LEASED PREMISES ---
    add_paragraph(doc, "The following described Leased Premises.",
                  space_before=6, space_after=6)

    # "Leased Premises" bold, then address with blue markers
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(0)
    run = p.add_run("Leased Premises")
    _apply_run_fmt(run, bold=True)
    add_mixed_runs(p, f"   #{B(building_num)} Sanctuary Boulevard")

    add_paragraph(doc, f"\t\t\tFloor {B(floor)}, Suite {B(suite)}",
                  space_after=0)
    add_paragraph(doc, "\t\t\tMandeville, LA  70471", space_after=0)
    add_paragraph(doc, f"\t\t\t{B(sqft)} sqft Rentable Area", space_after=6)

    # In Solido Liability
    add_clause(doc, "In Solido Liability",
        "If the above described Leased Premises is leased to more than one party, "
        "the obligations of all such parties hereunder, as lessees, shall be in solido.")

    # Term
    add_clause(doc, "Term",
        f"The primary term of this lease shall be {B(term_months)} months, commencing on "
        f"{B(format_date_long(start_date))} and ending on {B(format_date_long(end_date))}.")

    # Renewal
    add_clause(doc, "Renewal",
        f"The lease agreement will automatically renew for {B(renewal_months)} months at the "
        "expiration of the primary term and all renewal terms of Lessee.  Lessee's right to "
        "terminate the term of this Lease is expressly conditioned upon Lessee delivering to "
        f"Lessor written notice of exercise of such option not later than {B(notice_words)} "
        f"({B(notice_days)}) days prior to the expiration date of the primary term or the then "
        "renewal Term.  All of the terms and conditions of this Lease shall remain in full force "
        "and effect during the renewal term except that there will be no further option or right "
        "to renew or extend the term of this Lease.")

    # Rental and Place of Payment
    first_payment_month = first_payment.strftime("%B %Y")
    add_clause(doc, "Rental and Place of Payment",
        f"The rental under this lease shall be {B(rent_words)} monthly, "
        "payable in advance.  Rent for the first full calendar month of the term of this "
        "lease, plus the rent for any fractional month preceding such first calendar "
        "month, shall be payable on the signing of this lease by Lessee and rent for subsequent "
        f"months shall be payable on the first day of {B(first_payment_month)} and on the "
        "first day of each calendar month thereafter.  All payments of rent shall be made to "
        "Lessor at #1 Sanctuary Boulevard, Mandeville, Louisiana 70471, but Lessor may from "
        "time to time, with the written notice to Lessee, designate other persons and places "
        "for payment of rent.  Commencing with the first month of the second calendar year "
        f"of the lease term, the rental shall increase {B(escalation_words)} percent "
        f"({B(escalation)}%) over the previous year's rent.  This increase shall be applicable "
        "to all subsequent years including any renewal leases.")

    # Utility & Common Area Charges
    if utilities == "Lessor" and common_area == "None":
        utility_text = (f"All utility charges to be paid by {B('Lessor')}.  There are no "
                        "common area charge or real estate taxes to be paid by the Lessee.")
    elif utilities == "Lessor" and common_area == "Lessee Pays":
        utility_text = (f"All utility charges to be paid by {B('Lessor')}.  "
                        f"{B('Common area charges and real estate taxes shall be paid by the Lessee')}.")
    elif utilities == "Lessee" and common_area == "None":
        utility_text = (f"All utility charges to be paid by {B('Lessee')}.  There are no "
                        "common area charge or real estate taxes to be paid by the Lessee.")
    else:
        utility_text = f"All utility charges and common area charges to be paid by {B('Lessee')}."

    add_clause(doc, "Utility & Common Area Charges", utility_text)

    # Use of Premises - split across paragraphs like the original
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(0)
    run = p.add_run("Use of Premises")
    _apply_run_fmt(run, bold=True)
    run2 = p.add_run("  The leased premises shall be used only for the following purposes:")
    _apply_run_fmt(run2)

    # Use type on its own line (blue, indented)
    add_paragraph(doc, f"      {B(use_of_premises)}", space_before=4, space_after=4)

    # Rest of Use of Premises clause
    add_paragraph(doc,
        "The leased premises shall not be used for any unlawful purpose or in any manner "
        "that may damage or depreciate the same.  Lessee shall comply with all requirements "
        "of State, Parish, Municipal, Federal, and other public authorities, relating to the "
        "use and occupancy of the leased premises.  No auction sales, or other sales not in "
        "the ordinary course of Lessee's business, shall be conducted on the leased premises, "
        "without the prior written consent of Lessor.  Lessee shall comply with all of the "
        "reasonable rules and regulations pertaining to the building and businesses operated "
        "therein promulgated by Lessor from time to time.  Lessor may amend, modify, delete, "
        "or add new and additional reasonable rules and regulations for the use and care of "
        "the Premises, the Building or the Common Area.  In the event of any breach of any "
        "rules and regulations set forth or any amendments or additions thereto, Lessor shall "
        "have all remedies in this Lease provided for default of Lessee.",
        space_before=0, space_after=4)

    # Build Out Specifications
    add_clause(doc, "Build Out Specifications", B(buildout))

    # Repairs and Maintenance
    add_clause(doc, "Repairs and Maintenance",
        "The leased premises and all appurtenances contained therein, including, but not "
        "limited to, fixtures, locks, keys, glass, elevators, plumbing, heating equipment "
        "and air conditioning equipment are accepted by Lessee in their present condition, "
        "including any vices or defects, latent or otherwise, that may now exist or hereafter "
        "arise in the leased premises, except as to such repairs or improvements as this lease "
        "requires Lessor to make.  Lessor shall maintain the following items of the leased "
        "premises in good order and repair, but shall not be required to make any other "
        "repairs or replacements whatsoever to the leased premises, except those rendered "
        "necessary by fire or other perils which would be covered by fire and extended "
        "coverage insurance:")

    add_sub_items(doc, [
        "Roof",
        "Elevator",
        "Building Exterior and Parking Areas",
        "HVAC and Plumbing Systems",
        "Electrical Systems Except Fixtures",
        "Structural Components",
        "Janitorial Services per the Sanctuary Office Park Cleaning Schedule."
    ])

    add_paragraph(doc,
        "Lessee shall, at Lessee's expense and within a reasonable period of time, make any "
        "and all repairs and replacements of whatsoever nature or character that may become "
        "necessary to the leased premises during the term of this lease other than those "
        "herein above required to be made by Lessor, including but not limited to:",
        space_before=4, space_after=2)

    add_sub_items(doc, [
        "All interior finish surfaces including floors, walls and ceilings.",
        "Light bulb and light fixture replacement.",
        "Interior window cleaning."
    ])

    add_paragraph(doc,
        "At the termination of this lease, Lessee shall return the leased premises to "
        "Lessor, in like order and condition as received, broom clean and free from trash, "
        "ordinary decay, wear and tear excepted, and shall deliver the keys to the leased "
        "premises to Lessor.", space_before=4, space_after=4)

    # Responsibility for Damages, Injuries, Losses
    add_clause(doc, "Responsibility for Damages, Injuries, Losses",
        "Lessor shall not be responsible for damage to lessee property or injury to persons "
        "or other losses or damages caused by or resulting from leaks in the roof or plumbing "
        "of the leased premises, unless Lessor fails to take steps toward repairing such leaks "
        "within a reasonable period of time after being notified thereof by Lessee.  Should "
        "Lessee fail to so notify Lessor promptly, Lessee shall be responsible for damages or "
        "losses resulting to Lessor or third parties.  Lessor shall not be responsible for "
        "damage to property or injury to person or other losses or damages caused by or "
        "resulting from vices or defects, latent or otherwise, that may now exist or hereafter "
        "arise in the leased premise, or caused by or resulting from disrepair, damage or "
        "conditions necessitating repairs or replacements required herein to be made by Lessee.")

    add_paragraph(doc,
        "Lessor indemnifies Lessee for all damages arising because of the negligence or "
        "willful misconduct of its agents or employees arising from its ownership of the "
        "building.  Lessee indemnifies Lessor for any negligent act or willful misconduct of "
        "its employees, agents, invitees on the premises or arising from the business Lessee "
        "transacts on the premises.  The parties hereby release each other and their respective "
        "employees, directors, agents or invitees from any and all liability or responsibility "
        "to the other by way of subrogation or otherwise for any loss or damage based on "
        "Lessee's occupancy of the premises or Lessor's ownership of the building, unless such "
        "loss is caused by the others negligence or willful misconduct.  The parties agree to "
        "notify their respective insurers of the release of subrogation claims.  These waivers "
        "are effective regardless of insurance coverage.",
        space_before=4, space_after=4)

    # Delayed Possession
    add_clause(doc, "Delayed Possession",
        "Should Lessor be delayed in delivering possession of the premises to Lessee on the "
        "commencement date of this lease, because of any delay of existing occupants to vacate "
        "or because of the construction of improvements or the making of repairs required by "
        "this lease to be made by Lessor not having been completed or because of any other "
        "reason, not due to the design of Lessor, this lease shall not be affected thereby and "
        "Lessee shall not be entitled to any damages for such delay, except that Lessee shall "
        "be allowed a remission of rent for the period prior to delivery of possession, in "
        "which case the termination date of this lease shall remain unchanged.")

    # Delay in Making Repairs
    add_clause(doc, "Delay in Making Repairs",
        "If this lease requires Lessor to make improvements or repairs to the leased premises "
        "and Lessor deems it impracticable to do so prior to the commencement date of this "
        "lease, Lessee agrees that Lessor may make such improvements or repairs after "
        "possession is delivered to Lessee, in a manner such as not to unreasonably interfere "
        "with the operation of Lessee's business, in which case there shall be no reduction "
        "or remission of rent.")

    # Alterations or Additions by Lessee
    add_clause(doc, "Alterations or Additions by Lessee",
        "Lessee shall not make any alterations or additions to the leased premises, without "
        "obtaining Lessor's prior written consent. Such approval by Lessor shall not be "
        "unreasonably withheld.  Any and all alterations, additions or other improvements made "
        "by Lessee, with or without the consent of Lessor regardless of how attached (except "
        "movable trade fixtures and equipment), shall become the property of Lessor upon "
        "termination of this lease, without compensation therefore to Lessee, unless a "
        "provision is made at the time of approval by Lessor to allow the removal by Lessee.  "
        "Lessor shall have the right to require that Lessee, prior to the termination of this "
        "lease, remove any or all such alterations, additions or improvements and restore the "
        "leased premises to their condition at the time of commencement of this lease.")

    # Insurance
    add_clause(doc, "Insurance",
        "Lessee shall provide and maintain, for the mutual benefit of Lessee and Lessor, "
        "liability insurance against claims (1) for bodily injury, or death resulting "
        "therefrom, occurring on the leased premises, in the amount of "
        f"{B(insurance_formatted)} as to any one occurrence, and (2) for property damage in "
        f"the amount of {B(insurance_formatted)} as to any one occurrence on the leased "
        "premises.")

    add_paragraph(doc,
        "Lessor shall be named as an insured in the policies providing "
        "such insurance and certificates of insurance, evidencing such insurance, shall be "
        "delivered to Lessor promptly upon the execution of this lease.  All of said insurance "
        "shall be carried with responsible insurance companies authorized to transact business "
        "in the State of Louisiana and shall not be canceled or materially altered by Lessee, "
        "without thirty (30) days prior written notice to Lessor.  Lessee assumes full "
        "responsibility for the condition of the leased premises and agrees to hold Lessor "
        "harmless from all liability for injury to persons or damage to property or other "
        "losses or damages caused by or resulting from any accident or other occurrence in, on "
        "or about the leased premises.  Lessee shall be responsible for insuring his or her "
        "own office contents.",
        space_before=0, space_after=4)

    # Acts of Lessee Affecting Insurance
    add_clause(doc, "Acts of Lessee Affecting Insurance",
        "If the rate of fire or other casualty insurance covering the leased premises is "
        "increased due to acts of Lessee, Lessee shall pay to Lessor the increased cost of "
        "such insurance.  Lessee will not do or cause or suffer to be done any act or thing "
        "whereby the policy or policies of fire or other casualty insurance covering the "
        "leased premises shall become void or suspended.  Should Lessee's occupancy cause "
        "Lessor to be unable to obtain fire or other casualty insurance covering the leased "
        "premises, Lessor shall have the right to terminate this lease upon giving Lessee not "
        "less than ten (10) days prior written notice.  Lessee agrees to notify Lessor at any "
        "time the leased premises will become unoccupied, so that Lessor may obtain necessary "
        "vacancy permits from Lessor's insurers.")

    # Signs by Lessee - lobby name on its own centered line
    add_clause(doc, "Signs by Lessee",
        "Lessee shall have the right to erect and maintain signs advertising Lessee's "
        "business on the exterior of the leased premises however only within the interior of "
        "the Lessee's floor of the building.  It is further provided that such signs shall "
        "conform to the size and design consistent with that expected for professional "
        "business offices and maintained in accordance with the rules and regulations of the "
        "properly constituted authorities.  Lessee shall remove all signs placed on the leased "
        "premises at the expiration of this lease and shall repair damage to the leased "
        "premises caused by the erection, maintenance or removal of such signs.  "
        "Lessee's company name shall be listed on the main lobby directory as follows:")

    # Lobby name on own centered line
    add_paragraph(doc, B(lobby_name),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)

    # Right of Entry by Lessor
    add_clause(doc, "Right of Entry by Lessor",
        "Lessor shall have the right to enter the leased premises at all reasonable times for "
        "the purpose of inspecting the same and for the purpose of making repairs "
        "required to be made by Lessor or which Lessor may deem necessary to make.")

    # For Sale and For Rent Signs; Inspection Prospects
    add_clause(doc, "For Sale and For Rent Signs; Inspection Prospects",
        'Lessor shall have the right to place the usual "For Sale" and "By Auction" signs on '
        "the leased premises at any time during the entire term of this lease and the usual "
        '"For Rent" signs on the leased premises during the last six (6) months of the term '
        "of this lease.  Lessee agrees to allow persons authorized by Lessor to inspect the "
        "leased premises during the entire term of this lease with the view of purchasing the "
        "same and during the last six (6) months of the term of this lease with the view of "
        "renting the same, such inspections to be at reasonable hours.")

    # Surrender of Possession
    add_clause(doc, "Surrender of Possession",
        "Upon expiration or termination of this lease, Lessee shall surrender possession of "
        "the leased premises immediately to Lessor.  Any holding over by Lessee shall not "
        "operate, except by written agreement, to extend or renew this lease, but in such "
        "case, Lessor may terminate Lessee's occupancy at once or may consider such occupancy "
        "to be from month to month; and Lessee, in the event of such holding over without "
        "Lessor's consent, shall pay double the rent stipulated in this lease, together with "
        "such loss or damage as may be caused to Lessor by such holding over.")

    # Subleasing or Assignment
    add_clause(doc, "Subleasing or Assignment",
        "Lessee shall not have the right to sublease the leased premises, in whole or in "
        "part, or to assign this lease or grant use of the leased premises to others, without "
        "the prior written consent of Lessor, provided that consent to sublease the leased "
        "premises in whole or to assign this lease shall not be unreasonably withheld.  Any "
        "such sublease shall contain all the provisions of this lease to the extent applicable.")

    # Damage by Fire or Other Casualty
    add_clause(doc, "Damage by Fire or Other Casualty",
        "If the leased premises are destroyed or damaged to an extent so as to render them "
        "wholly unfit for the purposes for which they are leased, by fire or other casualty, "
        "this lease shall automatically terminate, provided such destruction or damage is not "
        "caused by the neglect or design of Lessee.  If, however, the leased premises are "
        "damaged by fire or other casualty and can be repaired within one hundred twenty (120) "
        "days after the date of such fire or other casualty, this lease shall not terminate "
        "and Lessor shall give notice to Lessee, within thirty (30) days after such fire or "
        "other casualty, that Lessor will repair such damage, at Lessor's costs, within said "
        "one hundred and twenty (120) days period, in which case Lessee shall be entitled to a "
        "reduction or remission of rent such as shall be just and proportionate but shall not "
        "be entitled to any other damages; provided that if Lessor fails to complete such "
        "repairs within said one hundred and twenty (120) day period, because of causes not "
        "due to the fault or design of Lessor, this lease shall not terminate and Lessee "
        "shall not be entitled to damages, but shall be entitled only to further just and "
        "proportionate reduction or remission of rent.")

    # Default
    add_clause(doc, "Default",
        "If Lessee fails to pay any installment of rent due under this lease or fails to "
        "comply with any other provision of this lease, within ten (10) days after notice by "
        "Lessor to Lessee demanding same, provided that said notice need not be given with "
        "regard to nonpayment of rent after such notice has been given twice during the period "
        "of this lease, or if Lessee discontinues the use of the leased premises for the "
        "purposes for which leased, or abandons the leased premises, or removes from the "
        "leased premises any property against which Lessor is entitled to a lessor's lien or "
        "makes an assignment for the benefit of creditors or is adjudged a bankrupt in an "
        "involuntary bankruptcy proceeding, or files any type of proceeding or applies for any "
        "relief under the laws for the United States relating to bankruptcy or State Laws "
        "relating to insolvency, or if a receiver or other custodian is appointed for Lessee "
        "or any of Lessee's property by any court, then, in any such events, Lessor shall "
        "have the right, at Lessor's option, without putting Lessee in default and without "
        "notice of default, (1) to cancel this lease effective immediately or effective as of "
        "any date Lessor may select or (2) to proceed one or more times for past due "
        "installments of rent only, without prejudicing the right to proceed later or "
        "additional installments or exercise any other remedy, or (3) to declare the unpaid "
        "rent for the entire unexpired term of this lease immediately due and payable and at "
        "once demand and receive payment thereof or (4) to have recourse to any other remedy "
        "or mode of redress to which Lessor may be allowed by law.  In the event Lessor "
        "exercises the right to cancel this lease, then (a) Lessor shall have the right, as "
        "soon as said cancellation is effective, to re-enter the leased premises and re-let "
        "the same for such price and on such terms as may be immediately available, without "
        "prior written notice or court proceedings, Lessee hereby assenting thereto and "
        "expressly waiving any notice to vacate, and (b) Lessee shall be and remain liable "
        "not only for all rent payable to the date such cancellation becomes effective, but "
        "also for all damage or loss suffered by Lessor for the remaining term of this lease "
        "resulting from such cancellation.  Failure of Lessor to exercise any right granted in "
        "this paragraph shall not be construed as a waiver of the right to subsequently "
        "enforce for a new default such right and no indulgence by Lessor shall be construed "
        "as a waiver of any right herein granted.")

    # Attorney's Fees
    add_clause(doc, "Attorney's Fees",
        "Should an attorney be engaged by Lessor to enforce payment of the rent due under "
        "this lease or to protect any of the interests of Lessor hereunder, with or without "
        "judicial proceedings, Lessee agrees to pay Lessor the reasonable fee of such "
        "attorney, which fee is hereby fixed, if the collection of money is involved at 25% "
        "of the amount of such money, such fee in no event to be less than $100.00, and "
        "Lessee also agrees to pay all court costs and other expense incurred by Lessor.")

    # Release of Lessor on Sale
    add_clause(doc, "Release of Lessor on Sale",
        "Upon a sale or transfer of the leased premises, by Lessor or a subsequent purchaser "
        "or transferor thereof, the purchaser or transferee by virtue of such sale or transfer "
        "shall be bound for the performance of all of Lessor's agreements and obligations "
        "under this lease and the vendor or transferor shall thereupon be released from any "
        "and all liability thereafter arising under this lease.")

    # Notices
    add_clause(doc, "Notices",
        "Any notice to be given under this lease by Lessor to Lessee shall be considered as "
        "duly given, whether received or not, if made in writing, addressed to Lessee and "
        "mailed by registered or certified mail to Lessee at the leased premises.  Any notice "
        "to be given under this lease by Lessee to Lessor shall be considered as duly given, "
        "whether received or not, if made in writing, addressed to Lessor and mailed by "
        "registered or certified mail to Lessor at the place where rent is required to be "
        "paid under this lease as above provided.  Either Lessor or Lessee may change the "
        "designated place to which written notice is to be sent, by so advising the other, in "
        "writing, by registered or certified mail, at the place designated in this lease or "
        "such place as may have been subsequently designated in accordance with this paragraph.")

    # Security Deposit
    sec_dep = security_deposit.strip()
    if sec_dep.lower() == "waived":
        sec_text = B("Waived")
    else:
        sec_text = B(format_currency(sec_dep))

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(2)
    run = p.add_run("Security Deposit")
    _apply_run_fmt(run, bold=True)
    run2 = p.add_run("   ")
    _apply_run_fmt(run2)
    add_mixed_runs(p, sec_text)

    # Waiver (no title in original)
    add_paragraph(doc,
        "Failure of Lessor to require a strict performance by Lessee of any of the covenants, "
        "provisions, or conditions of this lease, on one or more occasions, shall not "
        "constitute a waiver by Lessor of the right thereafter to require strict compliance "
        "with said covenants, provisions, and conditions.",
        space_before=6, space_after=4)

    # Governing Law (no title in original)
    add_paragraph(doc,
        "This lease shall be deemed to be a contract made under the laws of the State of "
        "Louisiana and shall be construed in accordance with and governed by the laws of the "
        "State of Louisiana and ordinances of the municipality and parish where the leased "
        "premises are situated and the rules and regulations of their duly constituted "
        "authorities.",
        space_before=6, space_after=4)

    # Binding Effect (no title in original)
    add_paragraph(doc,
        "All the provisions contained herein shall be binding upon and shall ensure the "
        "benefit of the parties hereto their heirs, executors, administrators, successors, "
        "and assigns.",
        space_before=6, space_after=4)

    # Entire Agreement (no title in original)
    add_paragraph(doc,
        "The whole agreement between the parties hereto is set forth in this instrument and "
        "they shall not be bound by any agreements, conditions, understandings, or "
        "representations other than are expressly stipulated and set forth herein or in any "
        "amendments hereto.",
        space_before=6, space_after=4)

    # Attachments (no title in original)
    add_paragraph(doc,
        "See attached Floor plan, Janitorial Duty List and Sanctuary Office Park Rules "
        "and Regulations.",
        space_before=6, space_after=4)

    # --- SIGNATURE BLOCKS ---
    add_paragraph(doc, "", space_before=24, space_after=12)

    table = doc.add_table(rows=6, cols=2)
    table.allow_autofit = False
    _remove_table_borders(table)

    for row_obj in table.rows:
        row_obj.cells[0].width = Inches(3.25)
        row_obj.cells[1].width = Inches(3.25)

    _sig_cell(table.cell(0, 0), "LESSEE:", bold=True)
    _sig_cell(table.cell(0, 1), "LESSOR:", bold=True)

    _sig_cell(table.cell(1, 0), fields['lessee_name'], color=BLUE)
    _sig_cell(table.cell(1, 1), "Crosby Development Company, LLC")

    # Spacer row for handwritten signature
    _sig_cell(table.cell(2, 0), "")
    _sig_cell(table.cell(2, 1), "")
    table.rows[2].height = Inches(0.5)

    _sig_cell(table.cell(3, 0), "By: ___________________________")
    _sig_cell(table.cell(3, 1), "By: ___________________________")

    _sig_cell(table.cell(4, 0), "Title: ________________________")
    _sig_cell(table.cell(4, 1), "Title: ________________________")

    _sig_cell(table.cell(5, 0), "Date: _________________________")
    _sig_cell(table.cell(5, 1), "Date: _________________________")

    # --- IN SOLIDO OBLIGATION (before other appendices) ---
    in_solido = fields.get("in_solido", "").strip().lower()
    if in_solido in ("1", "true", "yes"):
        doc.add_page_break()
        add_in_solido_appendix(doc)

    # --- APPENDIX A: JANITORIAL DUTY LIST ---
    doc.add_page_break()
    add_janitorial_appendix(doc)

    # --- APPENDIX B: FLOOR PLAN ---
    floorplan_path = fields.get("floorplan_path", "").strip()
    if floorplan_path and os.path.isfile(floorplan_path):
        doc.add_page_break()
        add_floorplan_appendix(doc, floorplan_path)

    return doc, lease_id


# ---------------------------------------------------------------------------
# Appendices
# ---------------------------------------------------------------------------

def add_in_solido_appendix(doc):
    """Append the In Solido Obligation and Guaranty addendum."""
    add_paragraph(doc, "IN SOLIDO OBLIGATION AND GUARANTY", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14,
                  space_before=0, space_after=16)

    add_paragraph(doc,
        "For value received and to induce the Lessor to enter into the foregoing "
        "lease, the undersigned hereby makes himself or itself a party to said lease "
        "and binds himself or itself in solido with the Lessee or Lessees under said "
        'lease (hereinafter referred to as "Lessee") for the faithful performance and '
        "fulfillment by Lessee of all of Lessee's agreements and obligations contained "
        "in said lease and guarantees to Lessor and Lessor's heirs, executors, "
        "administrators, successors and assigns, the punctual payment of all rents due "
        "under said lease and the performance of all other agreements and obligations "
        "of Lessee contained in said lease; the undersigned consenting to extensions "
        "of payment of rent by Lessor and other indulgences by Lessor to Lessee and "
        "amendments and modifications entered into between Lessor and Lessee regarding "
        "said lease, and waiving any and all requirements of notice demand, "
        "non-payment, non-performance or dishonor and all other requirements of law.",
        space_before=0, space_after=16, first_line_indent=0.5)

    # Two signature blocks for In Solido Obligors
    for _ in range(2):
        table = doc.add_table(rows=5, cols=2)
        table.allow_autofit = False
        _remove_table_borders(table)

        for row_obj in table.rows:
            row_obj.cells[0].width = Inches(2.5)
            row_obj.cells[1].width = Inches(4.0)

        _sig_cell(table.cell(0, 0), "Dated______________________")
        _sig_cell(table.cell(0, 1), "________________________________")

        # Label under first signature line
        p = table.cell(0, 1).paragraphs[0]
        p.clear()
        run = p.add_run("________________________________")
        _apply_run_fmt(run, font_size=12)

        _sig_cell(table.cell(1, 0), "")
        p = table.cell(1, 1).paragraphs[0]
        p.clear()
        run = p.add_run("In Solido Obligor and Guarantor")
        _apply_run_fmt(run, font_size=10, italic=True)

        _sig_cell(table.cell(2, 0), "")
        p = table.cell(2, 1).paragraphs[0]
        p.clear()
        run = p.add_run("________________________________")
        _apply_run_fmt(run, font_size=12)
        p2 = table.cell(2, 1).add_paragraph()
        run2 = p2.add_run("Name Printed")
        _apply_run_fmt(run2, font_size=10, italic=True)

        _sig_cell(table.cell(3, 0), "")
        p = table.cell(3, 1).paragraphs[0]
        p.clear()
        run = p.add_run("________________________________")
        _apply_run_fmt(run, font_size=12)
        p2 = table.cell(3, 1).add_paragraph()
        run2 = p2.add_run("Address")
        _apply_run_fmt(run2, font_size=10, italic=True)

        _sig_cell(table.cell(4, 0), "")
        p = table.cell(4, 1).paragraphs[0]
        p.clear()
        run = p.add_run("________________________________")
        _apply_run_fmt(run, font_size=12)
        p2 = table.cell(4, 1).add_paragraph()
        run2 = p2.add_run("City, State, Zip")
        _apply_run_fmt(run2, font_size=10, italic=True)

        # Spacing between the two blocks
        add_paragraph(doc, "", space_before=12, space_after=12)


def add_janitorial_appendix(doc):
    """Append the Sanctuary Office Park Janitorial Duty List as an appendix."""
    add_paragraph(doc, "APPENDIX A", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14,
                  space_before=0, space_after=2)
    add_paragraph(doc, "Sanctuary Office Park", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=13,
                  space_before=0, space_after=0)
    add_paragraph(doc, "Janitorial Service - Duty List",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=12,
                  space_before=0, space_after=12)

    add_paragraph(doc, "Janitorial Duties", bold=True,
                  font_size=12, space_before=6, space_after=6)

    # Daily
    add_paragraph(doc, "Daily:", bold=True,
                  font_size=12, space_before=6, space_after=2)
    for item in [
        "Bathroom and Kitchen Trash Emptied. (SEE NOTE)",
        "Clean Toilets, Sinks and Vanities in Bathrooms",
        "Clean Kitchen Counters and Sinks (Dishes, Appliances and Tables "
            "are the responsibility of Tennant.)",
        "Sweep Lobbies",
        "Restock Paper Products",
    ]:
        add_paragraph(doc, f"    \u2022  {item}", font_size=11,
                      space_before=1, space_after=1)

    # Tuesday, Thursday and Friday
    add_paragraph(doc, "Tuesday, Thursday and Friday:", bold=True,
                  font_size=12, space_before=6, space_after=2)
    add_paragraph(doc, "    \u2022  Empty all office Trash (SEE NOTE)",
                  font_size=11, space_before=1, space_after=1)

    # Weekly
    add_paragraph(doc, "Weekly: (Done on Weekends)", bold=True,
                  font_size=12, space_before=6, space_after=2)
    for item in [
        "Fully Vacuum Carpets",
        "Sweep and Mop Tile",
        "Remove smudges from doors and door trim",
        "Clean walls around elevator push buttons",
        "Sweep down stairs",
    ]:
        add_paragraph(doc, f"    \u2022  {item}", font_size=11,
                      space_before=1, space_after=1)

    # Monthly
    add_paragraph(doc, "Monthly:", bold=True,
                  font_size=12, space_before=6, space_after=2)
    for item in [
        "Sweep Baseboards",
        "Dust Window Sills",
        "Dust Common Area (Building Lobbies) Furniture",
        "Wipe Down Kitchen Cabinets",
        "Change Office Can liners",
        "Clean fluorescent light lenses",
    ]:
        add_paragraph(doc, f"    \u2022  {item}", font_size=11,
                      space_before=1, space_after=1)

    # Quarterly
    add_paragraph(doc, "Quarterly:", bold=True,
                  font_size=12, space_before=6, space_after=2)
    for item in [
        "Change Air Filters",
        "Fill Bathroom Floor Drains",
        "Dust and check emergency lights",
    ]:
        add_paragraph(doc, f"    \u2022  {item}", font_size=11,
                      space_before=1, space_after=1)

    # Note
    note_text = (
        "Note: All food and drink items must be disposed of in kitchen trash "
        "containers. Liquids must be poured in the sink before it is placed in "
        "kitchen trash containers. Trash collection bags leak and will cause "
        "carpets to be stained. Carpet cleaning is the responsibility of the "
        "tenant. Cleaning of Microwaves, refrigerators, kitchen tables and "
        "office furniture are the responsibility of the tenant. If you have "
        "questions or issues, please contact the office @ 985-674-7500 or "
        "Charli @ 985-789-3784 | charli.cornell@gmail.com. For emergencies "
        "call 911 or Thomas Crosby @ 985-789-6252 | tommy@crosbydevelopment.com "
        "or Ryan Crosby @ 985-373-2383 | ryan@crosbydevelopment.com."
    )
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    run_bold = p.add_run("Note: ")
    _apply_run_fmt(run_bold, font_size=10, bold=True)
    run_body = p.add_run(note_text[6:])  # skip "Note: " prefix
    _apply_run_fmt(run_body, font_size=10)


def add_floorplan_appendix(doc, filepath):
    """Append a floor plan image or PDF as an appendix."""
    add_paragraph(doc, "APPENDIX B", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14,
                  space_before=0, space_after=2)
    add_paragraph(doc, "Floor Plan",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=13,
                  bold=True, space_before=0, space_after=12)

    ext = os.path.splitext(filepath)[1].lower()

    if ext in ('.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff', '.tif'):
        doc.add_picture(filepath, width=Inches(6.0))
        # Center the image
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif ext == '.pdf':
        try:
            from pdf2image import convert_from_path
            import sys
            poppler_path = os.path.join(sys._MEIPASS, 'poppler') if getattr(sys, 'frozen', False) else None
            images = convert_from_path(filepath, dpi=200, poppler_path=poppler_path)
            import tempfile
            for i, img in enumerate(images):
                tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
                img.save(tmp.name, 'PNG')
                tmp.close()
                if i > 0:
                    doc.add_page_break()
                doc.add_picture(tmp.name, width=Inches(6.0))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                os.unlink(tmp.name)
        except ImportError:
            add_paragraph(doc,
                "[ Floor plan PDF could not be embedded. "
                "Install pdf2image and poppler to enable PDF embedding. ]",
                italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=12, space_after=12)
    else:
        add_paragraph(doc,
            f"[ Unsupported file format: {ext} ]",
            italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=12, space_after=12)


# ---------------------------------------------------------------------------
# Amendment generation
# ---------------------------------------------------------------------------

def generate_amendment(fields):
    """Generate a lease amendment .docx from the field dictionary."""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    amd_id = get_next_lease_id("A")

    amd_num = fields["amendment_num"]
    lessee_name = fields["lessee_name"]
    suite = fields["suite"]
    amd_date = parse_date(fields["amendment_date"])
    orig_lease_date = fields["original_lease_date"]

    # --- AMENDMENT ID ---
    add_paragraph(doc, f"Document ID: {amd_id}",
                  alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_size=10,
                  space_before=0, space_after=4, italic=True)

    # --- HEADER ---
    add_paragraph(doc, f"AMENDMENT #{B(amd_num)} TO THE LEASE", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=13,
                  space_after=0)
    add_paragraph(doc,
        f'Between Crosby Development Company, LLC ("Lessor") and',
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_paragraph(doc, f'{B(lessee_name)} ("Lessee")',
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_paragraph(doc, B(suite),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_paragraph(doc, B(amd_date.strftime("%B %d, %Y").replace(" 0", " ")),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # --- LESSOR / LESSEE ---
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(0)
    run = p.add_run("Lessor:")
    _apply_run_fmt(run, bold=True)
    run2 = p.add_run("\tCrosby Development Company, LLC")
    _apply_run_fmt(run2)

    add_paragraph(doc, "\t\t1 Sanctuary Blvd.", space_after=0)
    add_paragraph(doc, "\t\tMandeville, LA  70471", space_after=8)

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(0)
    run = p.add_run("Lessee:")
    _apply_run_fmt(run, bold=True)
    run2 = p.add_run(f"\t{lessee_name}")
    _apply_run_fmt(run2, color=BLUE)

    lessee_addr = fields["lessee_address"]
    if lessee_addr.strip():
        for line in lessee_addr.strip().split("\n"):
            add_paragraph(doc, f"\t\t{line.strip()}", space_after=0, color=BLUE)
    add_paragraph(doc, "", space_after=4)

    # --- AMENDMENT BODY ---
    add_paragraph(doc,
        f"The parties hereby agree to amend the terms of the existing lease, dated "
        f"{B(orig_lease_date)}, as follows:",
        space_before=6, space_after=8)

    # Premises (optional)
    new_premises = fields.get("new_premises", "").strip()
    if new_premises:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(4)
        run = p.add_run("Premises:\t")
        _apply_run_fmt(run, bold=True)
        run2 = p.add_run(new_premises)
        _apply_run_fmt(run2, color=BLUE)

    # Rent
    rent = fields.get("new_rent", "").strip()
    if rent:
        rent_formatted = format_currency(rent)
        rent_words = dollars_to_words(rent)
        rent_date = fields.get("rent_effective_date", "").strip()
        if rent_date:
            rent_date_obj = parse_date(rent_date)
            rent_date_str = f"Beginning {format_date_long(rent_date_obj)}, t"
        else:
            rent_date_str = "T"

        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(4)
        run = p.add_run("Rent:\t\t")
        _apply_run_fmt(run, bold=True)
        run2 = p.add_run(f"{rent_date_str}he rent shall be {rent_words}.")
        _apply_run_fmt(run2, color=BLUE)

    # Improvements
    improvements = fields.get("improvements", "").strip()
    if improvements:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(4)
        run = p.add_run("Improvements:\t")
        _apply_run_fmt(run, bold=True)
        run2 = p.add_run(improvements)
        _apply_run_fmt(run2, color=BLUE)

    # Term
    new_end = fields.get("new_term_end", "").strip()
    if new_end:
        end_date_obj = parse_date(new_end)
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(4)
        run = p.add_run("Term:\t\t")
        _apply_run_fmt(run, bold=True)
        run2 = p.add_run(
            f"Shall be extended to {format_date_long(end_date_obj)}.")
        _apply_run_fmt(run2, color=BLUE)

    # Escalation
    escalation = fields.get("escalation", "").strip()
    if escalation:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(4)
        run = p.add_run("Escalation:\t")
        _apply_run_fmt(run, bold=True)
        run2 = p.add_run(
            f"The rent shall increase {escalation}% over the previous year's rent "
            "and applies to all renewals.")
        _apply_run_fmt(run2, color=BLUE)

    # Automatic Renewal
    renewal_months = fields.get("renewal_months", "").strip()
    notice_days = fields.get("notice_days", "").strip()
    if renewal_months:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(4)
        run = p.add_run("Automatic Renewal:\t")
        _apply_run_fmt(run, bold=True)
        run2 = p.add_run(
            f"This lease shall automatically renew for successive {renewal_months}-month "
            f"terms unless written notice is provided by either party at least "
            f"{notice_days} days in advance of expiration date to terminate the lease.")
        _apply_run_fmt(run2, color=BLUE)

    # Closing
    add_paragraph(doc, "", space_before=6, space_after=0)
    add_paragraph(doc,
        "All other terms and conditions of the original lease shall remain unchanged.",
        space_before=4, space_after=12)

    # --- SIGNATURE BLOCKS ---
    add_paragraph(doc, "", space_before=24, space_after=12)

    table = doc.add_table(rows=6, cols=2)
    table.allow_autofit = False
    _remove_table_borders(table)

    for row_obj in table.rows:
        row_obj.cells[0].width = Inches(3.25)
        row_obj.cells[1].width = Inches(3.25)

    _sig_cell(table.cell(0, 0), "LESSEE:", bold=True)
    _sig_cell(table.cell(0, 1), "LESSOR:", bold=True)

    _sig_cell(table.cell(1, 0), lessee_name, color=BLUE)
    _sig_cell(table.cell(1, 1), "Crosby Development Company, LLC")

    # Spacer row for handwritten signature
    _sig_cell(table.cell(2, 0), "")
    _sig_cell(table.cell(2, 1), "")
    table.rows[2].height = Inches(0.5)

    _sig_cell(table.cell(3, 0), "By: ___________________________")
    _sig_cell(table.cell(3, 1), "By: ___________________________")

    _sig_cell(table.cell(4, 0), "Title: ________________________")
    _sig_cell(table.cell(4, 1), "Title: ________________________")

    _sig_cell(table.cell(5, 0), "Date: _________________________")
    _sig_cell(table.cell(5, 1), "Date: _________________________")

    return doc, amd_id


# ---------------------------------------------------------------------------
# GUI
# ---------------------------------------------------------------------------

class LeaseGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Lease Generator - Crosby Development Company, LLC")
        self.root.geometry("720x900")
        self.root.resizable(True, True)

        # --- Save folder bar at top (shared across tabs) ---
        folder_frame = ttk.LabelFrame(root, text="Save Location")
        folder_frame.pack(fill="x", padx=10, pady=(10, 2))

        self.save_dir_var = tk.StringVar(value=get_save_dir())
        ttk.Label(folder_frame, text="Folder:").pack(
            side="left", padx=(8, 4), pady=6)
        save_entry = ttk.Entry(folder_frame, textvariable=self.save_dir_var,
                               width=55)
        save_entry.pack(side="left", fill="x", expand=True, padx=4, pady=6)
        ttk.Button(folder_frame, text="Browse...",
                   command=self._browse_save_dir).pack(
            side="right", padx=(4, 8), pady=6)

        notebook = ttk.Notebook(root)
        notebook.pack(fill="both", expand=True, padx=10, pady=(2, 10))

        lease_frame = ttk.Frame(notebook)
        notebook.add(lease_frame, text="  New Lease  ")
        self.build_lease_tab(lease_frame)

        amend_frame = ttk.Frame(notebook)
        notebook.add(amend_frame, text="  Amendment  ")
        self.build_amendment_tab(amend_frame)

    def _browse_save_dir(self):
        folder = filedialog.askdirectory(
            title="Choose Save Folder",
            initialdir=self.save_dir_var.get())
        if folder:
            self.save_dir_var.set(folder)
            cfg = load_config()
            cfg["save_dir"] = folder
            save_config(cfg)

    def build_lease_tab(self, parent):
        canvas = tk.Canvas(parent)
        scrollbar = ttk.Scrollbar(parent, orient="vertical", command=canvas.yview)
        scroll_frame = ttk.Frame(canvas)

        scroll_frame.bind("<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=scroll_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        row = 0
        self.lease_fields = {}

        def add_field(label, key, default="", width=50):
            nonlocal row
            ttk.Label(scroll_frame, text=label).grid(
                row=row, column=0, sticky="w", padx=(10, 5), pady=3)
            var = tk.StringVar(value=default)
            entry = ttk.Entry(scroll_frame, textvariable=var, width=width)
            entry.grid(row=row, column=1, sticky="ew", padx=(5, 10), pady=3)
            self.lease_fields[key] = var
            row += 1
            return entry

        def add_dropdown(label, key, options, default):
            nonlocal row
            ttk.Label(scroll_frame, text=label).grid(
                row=row, column=0, sticky="w", padx=(10, 5), pady=3)
            var = tk.StringVar(value=default)
            combo = ttk.Combobox(scroll_frame, textvariable=var, values=options,
                                 state="readonly", width=47)
            combo.grid(row=row, column=1, sticky="ew", padx=(5, 10), pady=3)
            self.lease_fields[key] = var
            row += 1

        def add_text(label, key, default="", height=3):
            nonlocal row
            ttk.Label(scroll_frame, text=label).grid(
                row=row, column=0, sticky="nw", padx=(10, 5), pady=3)
            text = tk.Text(scroll_frame, height=height, width=50,
                           font=("TkDefaultFont", 9))
            text.grid(row=row, column=1, sticky="ew", padx=(5, 10), pady=3)
            text.insert("1.0", default)
            self.lease_fields[key] = text
            row += 1

        ttk.Label(scroll_frame, text="LESSEE INFORMATION",
                  font=("TkDefaultFont", 10, "bold")).grid(
            row=row, column=0, columnspan=2, sticky="w", padx=10, pady=(10, 3))
        row += 1

        add_field("Lessee Name:", "lessee_name")
        add_field("Address Line 1:", "lessee_addr1")
        add_field("Address Line 2:", "lessee_addr2", "Mandeville, LA 70471")

        ttk.Label(scroll_frame, text="LEASED PREMISES",
                  font=("TkDefaultFont", 10, "bold")).grid(
            row=row, column=0, columnspan=2, sticky="w", padx=10, pady=(10, 3))
        row += 1

        add_dropdown("Building Number:", "building_num",
                     ["1", "2", "3", "4", "5"], "1")
        add_field("Floor:", "floor")
        add_field("Suite Number:", "suite")
        add_field("Square Footage:", "sqft")

        ttk.Label(scroll_frame, text="LEASE TERMS",
                  font=("TkDefaultFont", 10, "bold")).grid(
            row=row, column=0, columnspan=2, sticky="w", padx=10, pady=(10, 3))
        row += 1

        today = datetime.today()
        add_field("Lease Term (months):", "term_months", "36")
        add_field("Start Date (MM/DD/YYYY):", "start_date",
                  today.strftime("%m/%d/%Y"))

        default_end = today + relativedelta(months=36)
        add_field("End Date (MM/DD/YYYY):", "end_date",
                  default_end.strftime("%m/%d/%Y"))

        add_field("Renewal Period (months):", "renewal_months", "12")
        add_field("Notice Period (days):", "notice_days", "90")

        ttk.Label(scroll_frame, text="FINANCIAL",
                  font=("TkDefaultFont", 10, "bold")).grid(
            row=row, column=0, columnspan=2, sticky="w", padx=10, pady=(10, 3))
        row += 1

        add_field("Monthly Rent ($):", "rent")
        add_field("Rent Escalation (%):", "escalation", "3")

        # Default first payment = 1st of the month after lease start
        try:
            default_start = parse_date(self.lease_fields["start_date"].get())
        except (ValueError, KeyError):
            default_start = today
        first_pay = (default_start + relativedelta(months=1)).replace(day=1)
        add_field("First Rent Payment Date:", "first_payment_date",
                  first_pay.strftime("%m/%d/%Y"))

        add_dropdown("Utilities Paid By:", "utilities",
                     ["Lessor", "Lessee"], "Lessor")
        add_dropdown("Common Area Charges:", "common_area",
                     ["None", "Lessee Pays"], "None")

        ttk.Label(scroll_frame, text="PREMISES USE & DETAILS",
                  font=("TkDefaultFont", 10, "bold")).grid(
            row=row, column=0, columnspan=2, sticky="w", padx=10, pady=(10, 3))
        row += 1

        add_field("Use of Premises:", "use_of_premises", "General Office")
        add_text("Build Out Specs:", "buildout", "Space to remain as is.")
        add_field("Lobby Directory Name:", "lobby_name")
        add_field("Security Deposit ($):", "security_deposit", "500")
        add_field("Insurance Amount ($):", "insurance_amount", "1,000,000")

        # Floor Plan file picker
        ttk.Label(scroll_frame, text="APPENDICES",
                  font=("TkDefaultFont", 10, "bold")).grid(
            row=row, column=0, columnspan=2, sticky="w", padx=10, pady=(10, 3))
        row += 1

        ttk.Label(scroll_frame, text="Floor Plan File:").grid(
            row=row, column=0, sticky="w", padx=(10, 5), pady=3)
        fp_frame = ttk.Frame(scroll_frame)
        fp_frame.grid(row=row, column=1, sticky="ew", padx=(5, 10), pady=3)
        self.floorplan_var = tk.StringVar()
        self.lease_fields["floorplan_path"] = self.floorplan_var
        fp_entry = ttk.Entry(fp_frame, textvariable=self.floorplan_var, width=38)
        fp_entry.pack(side="left", fill="x", expand=True)
        fp_btn = ttk.Button(fp_frame, text="Browse...",
                            command=self._browse_floorplan)
        fp_btn.pack(side="right", padx=(5, 0))
        row += 1

        # In Solido toggle
        self.in_solido_var = tk.BooleanVar(value=False)
        self.lease_fields["in_solido"] = self.in_solido_var
        chk = ttk.Checkbutton(scroll_frame,
                               text="Include In Solido Obligation and Guaranty addendum",
                               variable=self.in_solido_var)
        chk.grid(row=row, column=0, columnspan=2, sticky="w", padx=10, pady=(6, 3))
        row += 1

        _recalc_guard = [False]

        def recalc_end(*_args):
            if _recalc_guard[0]:
                return
            try:
                months = int(self.lease_fields["term_months"].get())
                start = parse_date(self.lease_fields["start_date"].get())
                end = start + relativedelta(months=months)
                _recalc_guard[0] = True
                self.lease_fields["end_date"].set(end.strftime("%m/%d/%Y"))
            except (ValueError, KeyError):
                pass
            finally:
                _recalc_guard[0] = False

        def recalc_term(*_args):
            if _recalc_guard[0]:
                return
            try:
                start = parse_date(self.lease_fields["start_date"].get())
                end = parse_date(self.lease_fields["end_date"].get())
                diff = relativedelta(end, start)
                months = diff.years * 12 + diff.months
                if months > 0:
                    _recalc_guard[0] = True
                    self.lease_fields["term_months"].set(str(months))
            except (ValueError, KeyError):
                pass
            finally:
                _recalc_guard[0] = False

        def recalc_first_payment(*_args):
            try:
                start = parse_date(self.lease_fields["start_date"].get())
                fp = (start + relativedelta(months=1)).replace(day=1)
                self.lease_fields["first_payment_date"].set(
                    fp.strftime("%m/%d/%Y"))
            except (ValueError, KeyError):
                pass

        self.lease_fields["term_months"].trace_add("write", recalc_end)
        self.lease_fields["start_date"].trace_add("write", recalc_end)
        self.lease_fields["end_date"].trace_add("write", recalc_term)
        self.lease_fields["start_date"].trace_add("write", recalc_first_payment)

        def sync_lobby(*_args):
            lobby = self.lease_fields["lobby_name"]
            if not lobby.get().strip():
                lobby.set(self.lease_fields["lessee_name"].get())

        self.lease_fields["lessee_name"].trace_add("write", sync_lobby)

        row += 1
        btn = ttk.Button(scroll_frame, text="Generate Lease",
                         command=self.generate_lease_doc)
        btn.grid(row=row, column=0, columnspan=2, pady=20)

        scroll_frame.columnconfigure(1, weight=1)

    def build_amendment_tab(self, parent):
        canvas = tk.Canvas(parent)
        scrollbar = ttk.Scrollbar(parent, orient="vertical", command=canvas.yview)
        scroll_frame = ttk.Frame(canvas)

        scroll_frame.bind("<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=scroll_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        row = 0
        self.amend_fields = {}

        def add_field(label, key, default="", width=50):
            nonlocal row
            ttk.Label(scroll_frame, text=label).grid(
                row=row, column=0, sticky="w", padx=(10, 5), pady=3)
            var = tk.StringVar(value=default)
            entry = ttk.Entry(scroll_frame, textvariable=var, width=width)
            entry.grid(row=row, column=1, sticky="ew", padx=(5, 10), pady=3)
            self.amend_fields[key] = var
            row += 1

        def add_text(label, key, default="", height=3):
            nonlocal row
            ttk.Label(scroll_frame, text=label).grid(
                row=row, column=0, sticky="nw", padx=(10, 5), pady=3)
            text = tk.Text(scroll_frame, height=height, width=50,
                           font=("TkDefaultFont", 9))
            text.grid(row=row, column=1, sticky="ew", padx=(5, 10), pady=3)
            text.insert("1.0", default)
            self.amend_fields[key] = text
            row += 1

        today = datetime.today()

        ttk.Label(scroll_frame, text="AMENDMENT DETAILS",
                  font=("TkDefaultFont", 10, "bold")).grid(
            row=row, column=0, columnspan=2, sticky="w", padx=10, pady=(10, 3))
        row += 1

        add_field("Amendment Number:", "amendment_num")
        add_field("Lessee Name:", "lessee_name")
        add_text("Lessee Address:", "lessee_address", "", height=2)
        add_field("Suite / Premises:", "suite")
        add_field("Original Lease Date:", "original_lease_date")
        add_field("Amendment Date (MM/DD/YYYY):", "amendment_date",
                  today.strftime("%m/%d/%Y"))

        ttk.Label(scroll_frame, text="AMENDED TERMS",
                  font=("TkDefaultFont", 10, "bold")).grid(
            row=row, column=0, columnspan=2, sticky="w", padx=10, pady=(10, 3))
        row += 1

        add_text("New Premises Description:", "new_premises", "", height=2)
        add_field("New Rent Amount ($):", "new_rent")
        add_field("Rent Effective Date (MM/DD/YYYY):", "rent_effective_date")
        add_field("Improvements:", "improvements", "Space leases as is.")
        add_field("New Term End Date (MM/DD/YYYY):", "new_term_end")
        add_field("Escalation (%):", "escalation", "3")
        add_field("Auto Renewal Period (months):", "renewal_months", "12")
        add_field("Auto Renewal Notice (days):", "notice_days", "90")

        row += 1
        btn = ttk.Button(scroll_frame, text="Generate Amendment",
                         command=self.generate_amendment_doc)
        btn.grid(row=row, column=0, columnspan=2, pady=20)

        scroll_frame.columnconfigure(1, weight=1)

    def _get_field_value(self, fields_dict, key):
        """Get value from a StringVar, BooleanVar, or Text widget."""
        widget = fields_dict[key]
        if isinstance(widget, tk.BooleanVar):
            return str(widget.get())
        if isinstance(widget, tk.Text):
            return widget.get("1.0", "end").strip()
        return widget.get().strip()

    def _browse_floorplan(self):
        filepath = filedialog.askopenfilename(
            title="Select Floor Plan",
            filetypes=[
                ("Image & PDF files", "*.png *.jpg *.jpeg *.bmp *.gif *.tiff *.tif *.pdf"),
                ("Image files", "*.png *.jpg *.jpeg *.bmp *.gif *.tiff *.tif"),
                ("PDF files", "*.pdf"),
                ("All files", "*.*"),
            ]
        )
        if filepath:
            self.floorplan_var.set(filepath)

    def generate_lease_doc(self):
        required = ["lessee_name", "lessee_addr1", "floor", "suite", "sqft",
                     "rent", "start_date", "end_date"]
        for key in required:
            val = self._get_field_value(self.lease_fields, key)
            if not val:
                label = key.replace("_", " ").title()
                messagebox.showerror("Missing Field",
                    f"Please fill in: {label}")
                return

        fields = {}
        for key in self.lease_fields:
            fields[key] = self._get_field_value(self.lease_fields, key)

        try:
            doc, lease_id = generate_lease(fields)
        except Exception as e:
            messagebox.showerror("Error", f"Error generating lease:\n{e}")
            return

        save_dir = self.save_dir_var.get()
        os.makedirs(save_dir, exist_ok=True)
        lessee = fields["lessee_name"]
        default_name = f"{lessee} - Lease ({lease_id}).docx"
        filepath = os.path.join(save_dir, default_name)
        doc.save(filepath)
        messagebox.showinfo("Success",
            f"Lease saved to:\n{filepath}\n\nLease ID: {lease_id}")

    def generate_amendment_doc(self):
        required = ["amendment_num", "lessee_name", "suite",
                     "original_lease_date", "amendment_date"]
        for key in required:
            val = self._get_field_value(self.amend_fields, key)
            if not val:
                label = key.replace("_", " ").title()
                messagebox.showerror("Missing Field",
                    f"Please fill in: {label}")
                return

        fields = {}
        for key in self.amend_fields:
            fields[key] = self._get_field_value(self.amend_fields, key)

        try:
            doc, amd_id = generate_amendment(fields)
        except Exception as e:
            messagebox.showerror("Error", f"Error generating amendment:\n{e}")
            return

        save_dir = self.save_dir_var.get()
        os.makedirs(save_dir, exist_ok=True)
        lessee = fields["lessee_name"]
        amd_num = fields["amendment_num"]
        default_name = f"{lessee} - Amendment #{amd_num} ({amd_id}).docx"
        filepath = os.path.join(save_dir, default_name)
        doc.save(filepath)
        messagebox.showinfo("Success",
            f"Amendment saved to:\n{filepath}\n\nDocument ID: {amd_id}")


def main():
    root = tk.Tk()
    app = LeaseGeneratorApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
